# Importing BeautifulSoup class from the bs4 module
from __future__ import print_function
from bs4 import BeautifulSoup
import os.path
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from google.oauth2 import service_account
from docx.api import Document
import docx
from docx.enum.text import WD_COLOR_INDEX
from termcolor import colored

SHOW_EVERYTHING_INCLUDING_NON_HIGHLIGHTED = False

# If modifying these scopes, delete the file token.json.

# The ID and range of a sample spreadsheet.
SAMPLE_SPREADSHEET_ID = '1-CZJljjGtjP9J2AX5n80jlOdPWBCvum5pm7F8rDS510'

service = build('sheets', 'v4', credentials=creds)

        # Call the Sheets API
sheet = service.spreadsheets()
#result = sheet.values().get(spreadsheetId=SAMPLE_SPREADSHEET_ID,
#                                   range="GRADES!A1:G13").execute()
#values = result.get('values', [])

SCOPES = ['https://www.googleapis.com/auth/spreadsheets']
SERVICE_ACCOUNT_FILE = 'keys.json'
creds = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE, scopes=SCOPES)

# for outputting pretty colors to the terminal
def convert_wd_color_index_to_termcolor(color_index):
    if (color_index == WD_COLOR_INDEX.BLUE):
        return "blue"
    if (color_index == WD_COLOR_INDEX.TURQUOISE):
        return "cyan"
    if (color_index == WD_COLOR_INDEX.GREEN):
        return "green"
    if (color_index == WD_COLOR_INDEX.BRIGHT_GREEN):
        return "light_green"
    if (color_index == WD_COLOR_INDEX.RED):
        return "red"
    if (color_index == WD_COLOR_INDEX.YELLOW):
        return "yellow"

    print("WARNING: Unrecognized color index: %s" % (color_index))
    return "white"


def calculateScoreFromHighlights(highlights):
    score = 0
    for h in highlights:
        score += h[1]
    return score


document = docx.Document('test2.docx')

print(colored("\n========= Found %d tables in the document ==========" % (len(document.tables)), "blue"))

# Print out summary of tables found in the document
for t, table in enumerate(document.tables):
    print(colored("Table %d has %d rows and %d columns" % (t, len(table.rows), len(table.columns)), "yellow"))

# each of these lists will contain tuples of (text, score) which we'll later remove dupes using set
highlightedFoundationals = []
highlightedProficients = []
highlightedExemplarys = []

# Print out detailed contents of each table, along with what is highlighted
for t, table in enumerate(document.tables):
    print("\n\nTABLE %d:" % (t))
    for r, row in enumerate(table.rows):
        print("\n-------- Row %d --------" % (r))
        for c, cell in enumerate(row.cells):
            print()
            for p, paragraph in enumerate(cell.paragraphs):
                if (len(paragraph.runs) == 0):
                    continue
                # score will be the precentage of runs inside this paragraph that are highlighted
                numHighlightedRuns = 0
                text = paragraph.text
                for r2, run in enumerate(paragraph.runs):
                    if len(paragraph.runs) == 1:
                        if c == 1:
                            if paragraph.runs[0].font.highlight_color == 10:
                                highlightedFoundationals.append((text, 1))
                            elif paragraph.runs[0].font.highlight_color == 4:
                                highlightedFoundationals.append((text, 0.5))
                        if c == 2:
                            if paragraph.runs[0].font.highlight_color == 10:
                                highlightedProficients.append((text, 1))
                            elif paragraph.runs[0].font.highlight_color == 4:
                                highlightedProficients.append((text, 0.5))
                        if c == 3:
                            if paragraph.runs[0].font.highlight_color == 10:
                                highlightedExemplarys.append((text, 1))
                            elif paragraph.runs[0].font.highlight_color == 4:
                                highlightedExemplarys.append((text, 0.5))
                    else:
                        colors_foundational = []
                        colors_proficient = []
                        colors_exemplary = []
                        if (c == 1):

                            for i in range(len(paragraph.runs)):
                                colors_foundational.append(paragraph.runs[i].font.highlight_color)

                        if (c == 2):
                            for i in range(len(paragraph.runs)):
                                colors_proficient.append(paragraph.runs[i].font.highlight_color)

                        if (c == 3):
                            for i in range(len(paragraph.runs)):
                                colors_exemplary.append(paragraph.runs[i].font.highlight_color)
                        colors_foundational = list(set(colors_foundational))
                        colors_proficient = list(set(colors_proficient))
                        colors_exemplary = list(set(colors_exemplary))

                        if len(colors_foundational) == 1:
                            if colors_foundational[0] == 10:
                                highlightedFoundationals.append((text, 1))
                            if colors_foundational[0] == 4:
                                highlightedFoundationals.append((text, .5))

                        elif len(colors_foundational) > 1:
                            if 10 and 4 in colors_foundational:
                                highlightedFoundationals.append((text, .75))
                            elif 10 in colors_foundational:
                                highlightedFoundationals.append((text, .5))
                            elif 4 in colors_foundational:
                                highlightedFoundationals.append([text, .25])

                        if len(colors_proficient) == 1:
                            if colors_proficient[0] == 10:
                                highlightedProficients.append((text, 1))
                            if colors_proficient[0] == 4:
                                highlightedProficients.append((text, .5))


                        elif len(colors_proficient) > 1:
                            if 10 and 4 in colors_proficient:
                                highlightedProficients.append((text, .75))
                            elif 10 in colors_proficient:
                                highlightedProficients.append((text, .5))
                            elif 4 in colors_proficient:
                                highlightedProficients.append((text, .25))

                        if len(colors_exemplary) == 1:
                            if colors_exemplary[0] == 10:
                                highlightedExemplarys.append((text, 1))
                            if colors_exemplary[0] == 4:
                                highlightedExemplarys.append((text, .5))


                        elif len(colors_exemplary) > 1:
                            if 10 and 4 in colors_exemplary:
                                highlightedExemplarys.append((text, .75))
                            elif 10 in colors_exemplary:
                                highlightedExemplarys.append((text, .5))
                            elif 4 in colors_exemplary:
                                highlightedExemplarys.append((text, .25))

                for r2, run in enumerate(paragraph.runs):
                    if run.font.highlight_color is not None:
                        print(
                            colored("*Table %d, Row %d, Cell %d, Paragraph %d, Run %d: %s" % (t, r, c, p, r2, run.text),
                                    convert_wd_color_index_to_termcolor(run.font.highlight_color)))
                    else:
                        if SHOW_EVERYTHING_INCLUDING_NON_HIGHLIGHTED:
                            print(" Table %d, Row %d, Cell %d, Paragraph %d, Run %d: %s" % (t, r, c, p, r2, run.text))

print("\n\n========= FOUNDATIONALS ==========")
highlightedFoundationals = list(set(highlightedFoundationals))
print(highlightedFoundationals)
print(calculateScoreFromHighlights(highlightedFoundationals))

print("\n\n========= PROFICIENTS ==========")
highlightedProficients = list(set(highlightedProficients))
# print(highlightedProficients)
print(calculateScoreFromHighlights(highlightedProficients))

print("\n\n========= EXEMPLARYS ==========")
highlightedExemplarys = list(set(highlightedExemplarys))
# print(highlightedExemplarys)
print(calculateScoreFromHighlights(highlightedExemplarys))


name = "test"
content_f = calculateScoreFromHighlights(highlightedFoundationals)
content_p = calculateScoreFromHighlights(highlightedProficients)
content_e = calculateScoreFromHighlights(highlightedExemplarys)
skills_f = 0
skills_p = 13
skills_e = 14
habits_f = 16
habits_p =18
habits_e =19
score = [[name,None,None,content_f,content_p,content_e,None,None,skills_f,skills_p,skills_e,None,None,habits_f,habits_p,habits_e]]
aoa = [["1/1/2020",4000]]
request = sheet.values().update(spreadsheetId=SAMPLE_SPREADSHEET_ID,
                                range="GRADES!B6",valueInputOption="USER_ENTERED",
                                body = {"values":score}).execute()


# Opening the html file
#HTMLFile = open("test.html", "r")

# Reading the file
#index = HTMLFile.read()

# soup reader
#soup = BeautifulSoup(index, 'html.parser')

# find all highlighted lines
#mydivs = soup.find_all("span", class_="c0 c13")
#score = [[len(mydivs)]]

#print(score)
