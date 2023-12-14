from __future__ import print_function
import sys
import subprocess

subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'pydrive','-q','--no-color'])
subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx','-q','--no-color'])
subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'docx','-q','--no-color'])
subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'termcolor','-q','--no-color'])
subprocess.check_call([sys.executable, '-m', 'pip', 'install', '--upgrade','google-api-python-client',
                       ' google-auth-httplib2','google-auth-oauthlib','-q','--no-color'])

from pydrive.auth import GoogleAuth
from pydrive.drive import GoogleDrive
import os
import os.path
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from google.oauth2 import service_account
from docx.api import Document
import docx
from docx.enum.text import WD_COLOR_INDEX
from termcolor import colored
from pathlib import Path

# *************************************************************** #
# Input variables

lightColor = 4
darkColor = 11

spreadsheetLink = '1-CZJljjGtjP9J2AX5n80jlOdPWBCvum5pm7F8rDS510'
spreadsheetNubricLink = '1zhUjd6sQ7vHY6O8BgEyjRIbO-U5UTRG-mrrfDENCVKI'
folderLink = '1tyeoX2ZgmgSW57cS8z-r4S_Z6hKPl2yo'

folderRoot = 'drive_download'

# *************************************************************** #


# *************************************************************** #
# Functions #

# Drive related

def authenticate():
    """
    Authenticates the loads google drive for folder input
    :params: none
    :return: authorized Google Drive
    """
    gauth = GoogleAuth()

    # Tries to load saved client credentials
    gauth.LoadCredentialsFile("credentials.txt")

    if gauth.credentials is None:
        # Authenticate if they're not there
        gauth.LocalWebserverAuth()
    elif gauth.access_token_expired:
        # Refresh them if expired
        gauth.Refresh()
    else:
        # Initialize the saved credentials
        gauth.Authorize()

    # Save the current credentials to a file
    gauth.SaveCredentialsFile("credentials.txt")

    drive = GoogleDrive(gauth)
    return drive

def escape_fname(name):
    """
    Simple function to reformat folder names
    :params: foldername
    :return: reformatted foldername
    """
    return name.replace('/','_')

def search_folder(folder_id, root):
    """
    Iterates through google drive folder until it finds google doc files to download
    If it finds relevant files, it downloads them to your local machine
    :param folder_id: URL ID of google drive folder
    :param root: computer root for download, i.e where the files should download
    :return: none, but downloads files
    """
    file_list = drive.ListFile({'q': "'%s' in parents and trashed=false" % folder_id}).GetList()
    for file in file_list:

        if file['mimeType'].split('.')[-1] == 'folder':
            foldername = escape_fname(file['title'])
            create_folder(root,foldername)
            search_folder(file['id'], '{}{}/'.format(root,foldername))
        else:
            download_mimetype = None
            filename = escape_fname(file['title'])
            filename = '{}{}'.format(root,filename)
            try:
                print('DOWNLOADING:', filename)
                if file['mimeType'] in MIMETYPES:
                    download_mimetype = MIMETYPES[file['mimeType']]

                    file.GetContentFile(filename+EXTENSIONS[file['mimeType']], mimetype=download_mimetype)
                else:
                    file.GetContentFile(filename)
            except:
                print('FAILED')
                f.write(filename+'\n')

def create_folder(path,name):
    """
    Creates local file to house downloaded templates
    :param name: name of created folder
    :param path: computer root for download, i.e where the files should download
    :return: none
    """
    os.mkdir('{}{}'.format(path,escape_fname(name)))


# Terminal related

# for outputting pretty colors to the terminal
def convert_wd_color_index_to_termcolor(color_index):
    """
    Converts color ID to string colors if recognized, else returns white
    :param name: name of created folder
    :param path: computer root for download, i.e where the files should download
    :return: none
    """
    if (color_index == WD_COLOR_INDEX.BLUE):
        return "blue"
    if (color_index == WD_COLOR_INDEX.TURQUOISE):
        return "cyan"
    if (color_index == WD_COLOR_INDEX.GREEN):
        return "green"
    if (color_index == WD_COLOR_INDEX.BRIGHT_GREEN):
        return "light_green"
    if (color_index == WD_COLOR_INDEX.RED):
        return "red"
    if (color_index == WD_COLOR_INDEX.YELLOW):
        return "yellow"

    print("WARNING: Unrecognized color index: %s" % (color_index))
    return "white"


def calculateScoreFromHighlights(highlights):
    """
    Calculates total numerical score by summing highlights
    :param highlights: array of all highlights and corresponding scores
    :return: numerical score
    """
    score = 0
    for h in highlights:
        score += h[1]
    return score

def convertToNubric(categoryArray, countTotalArray):
    """
    This function converts from a count of highlights to a 1-4 score for the Nubric
    :param categoryArray: array containing (foundational count, proficient count, exemplary count)
    :return: numerical score from 1-4 for the Nubric
    """
    nubricScore = ""

    if categoryArray[0] == countTotalArray[0]  and categoryArray[1] == countTotalArray[1] == 1 and categoryArray[2] >= 10:
        nubricScore = "4*"
    elif categoryArray[0] == countTotalArray[0] and categoryArray[1] == countTotalArray[1] and categoryArray[2] >= 3:
        nubricScore = "4"
    elif categoryArray[0] == countTotalArray[0] and categoryArray[1] >= (countTotalArray[1]* .7):
        nubricScore = "3"
    elif categoryArray[0] == countTotalArray[0]:
        nubricScore = "2"
    else:
        nubricScore = "1"

    return nubricScore


def count_total_in_range(document, table_num, start_row, end_row, start_column, end_column):

    allStandards = [0,0,0]
    # Print out detailed contents of each table, along with what is highlighted
    table = document.tables[table_num]
    # lists for counting within each table (has dupes)


    row_count = 0
    repeat_check = 0
    prior_row = None
    for r, row in enumerate(table.rows):
        this_row = row.cells[1].paragraphs[0].text
        if this_row != prior_row:
            prior_row = this_row
            row_count += 1
            # print("new row" + row.cells[1].paragraphs[0].text)

        if start_row <= row_count <= end_row:
            if row_count != repeat_check:
                repeat_check = row_count
                for c, cell in enumerate(row.cells):
                    if start_column <= c <= end_column:
                        for p, paragraph in enumerate(cell.paragraphs):
                            if len(paragraph.runs) == 0:
                                continue
                            text = paragraph.text
                            if (c == 1):  # 1st column is foundational
                                allStandards[0] = allStandards[0] + len(paragraph.runs)
                            if (c == 2):  # 2nd proficient
                                allStandards[1] = allStandards[1] + len(paragraph.runs)

                            if (c == 3):  # 3rd exemplary
                                allStandards[2] = allStandards[2] + len(paragraph.runs)

    return allStandards


def count_highlights_in_range(document, table_num, start_row, end_row, start_column, end_column, darkColor, lightColor):
    # each of these lists will contain tuples of (text, score) which we'll later remove dupes using set
    allFoundational = []
    allProficients = []
    allExemplarys = []

    # Print out detailed contents of each table, along with what is highlighted
    table = document.tables[table_num]
        # lists for counting within each table (has dupes)
    highlightedFoundationals = []
    highlightedProficients = []
    highlightedExemplarys = []

    row_count = 0
    repeat_check = 0
    prior_row = None
    for r, row in enumerate(table.rows):
        this_row = row.cells[1].paragraphs[0].text
        if this_row != prior_row:  
            prior_row = this_row
            row_count += 1
            #print("new row" + row.cells[1].paragraphs[0].text)
        
        if start_row <= row_count <= end_row:
            if row_count != repeat_check:
                repeat_check = row_count
                for c, cell in enumerate(row.cells):
                    if start_column <= c <= end_column:
                        for p, paragraph in enumerate(cell.paragraphs):
                            if len(paragraph.runs) == 0:
                                continue
                            text = paragraph.text
                            for r2, run in enumerate(paragraph.runs): #using runs to determine highlighted colors within the paragraphs. We found that when something is highlighted multiple colors, it will split into multiple runs
                                colors_foundational = []
                                colors_proficient = []
                                colors_exemplary = []
                                if (c == 1): #1st column is foundational
                                    for i in range(len(paragraph.runs)): #still within single paragraph. just getting all the different colors
                                        colors_foundational.append(paragraph.runs[i].font.highlight_color)  
                                                        
                                if (c == 2): #2nd proficient
                                    for i in range(len(paragraph.runs)):
                                        colors_proficient.append(paragraph.runs[i].font.highlight_color)

                                if (c == 3): #3rd exemplary
                                    for i in range(len(paragraph.runs)):
                                        colors_exemplary.append(paragraph.runs[i].font.highlight_color)
                                #set data type gets rid of all duplicates
                                colors_foundational = list(set(colors_foundational))
                                colors_proficient = list(set(colors_proficient))
                                colors_exemplary = list(set(colors_exemplary))
                                
                                #manually checking cases to determine score for the paragraph
                                if len(colors_foundational) == 1:
                                    # print(run.text)
                                    # print(colors_foundational)
                                    if colors_foundational[0] == darkColor:
                                        highlightedFoundationals.append((text,1))
                                    if colors_foundational[0] == lightColor:
                                        highlightedFoundationals.append((text,.5))
                                    
                                elif len(colors_foundational) > 1:
                                    # print(run.text)
                                    # print(colors_foundational)
                                    if darkColor in colors_foundational and lightColor in colors_foundational:
                                        highlightedFoundationals.append((text,.75))
                                    elif darkColor in colors_foundational:
                                        highlightedFoundationals.append((text,.5))
                                    elif lightColor in colors_foundational:
                                        highlightedFoundationals.append((text,.25))
                                

                                if len(colors_proficient) == 1:
                                    if colors_proficient[0] == darkColor:
                                        highlightedProficients.append((text,1))
                                    if colors_proficient[0] == lightColor:
                                        highlightedProficients.append((text,.5))
                                    
                                elif len(colors_proficient) > 1:
                                    if darkColor in colors_proficient and lightColor in colors_proficient:
                                        highlightedProficients.append((text,.75))
                                    elif darkColor in colors_proficient:
                                        highlightedProficients.append((text,.5))
                                    elif lightColor in colors_proficient:
                                        highlightedProficients.append((text,.25))

                                if len(colors_exemplary) == 1:
                                    if colors_exemplary[0] == darkColor:
                                        highlightedExemplarys.append((text,1))
                                    if colors_exemplary[0] == lightColor:
                                        highlightedExemplarys.append((text,.5))

                                    
                                elif len(colors_exemplary) > 1:
                                    if darkColor in colors_exemplary and lightColor in colors_exemplary:
                                        highlightedExemplarys.append((text,.75))
                                    elif darkColor in colors_exemplary:
                                        highlightedExemplarys.append((text,.5))
                                    elif lightColor in colors_exemplary:
                                        highlightedExemplarys.append((text,.25))


    # Calculate counts and return
    foundational_count = calculateScoreFromHighlights(list(highlightedFoundationals))
    proficient_count = calculateScoreFromHighlights(list(highlightedProficients))
    exemplary_count = calculateScoreFromHighlights(list(highlightedExemplarys))
    
    # print(allFoundational, allProficients, allExemplarys)

    return foundational_count, proficient_count, exemplary_count

# *************************************************************** #

SHOW_EVERYTHING_INCLUDING_NON_HIGHLIGHTED = False

# START OF DRIVE THINGS #

MIMETYPES = {
        # Drive Document files as MS dox
        'application/vnd.google-apps.document': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        # Drive Sheets files as MS Excel files.
        'application/vnd.google-apps.spreadsheet': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        # Drive presentation as MS pptx
        'application/vnd.google-apps.presentation': 'application/vnd.openxmlformats-officedocument.presentationml.presentation'
        # see https://developers.google.com/drive/v3/web/mime-types
    }
EXTENSIONS = {
        'application/vnd.google-apps.document': '.docx',
        'application/vnd.google-apps.spreadsheet': '.xlsx',
        'application/vnd.google-apps.presentation': '.pptx'
}

# creating the folder, downloading files from drive
if __name__ == '__main__':
    drive = authenticate()

    f = open("failed.txt","w+")
    folder_id = folderLink
    root = folderRoot
    if not os.path.exists(root):
        os.makedirs(root)

    search_folder(folder_id,root+'/')
    f.close() 


# END OF DRIVE THINGS #

dir_list = os.listdir(root)
# print(dir_list)


# If modifying these scopes, delete the file token.json.


SCOPES = ['https://www.googleapis.com/auth/spreadsheets']
SERVICE_ACCOUNT_FILE = 'keys.json'
creds = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE, scopes=SCOPES)
service = build('sheets', 'v4', credentials=creds)
sheet = service.spreadsheets()

# Start of doc things



# The big for loop which iterates through all the downloaded files in the folder. Count is used to help iterate through some 
# of the arrays.

count = 0

for file in Path(root).iterdir():
    name = os.listdir(root)[count][:-46] #46 charas is the "end year linear algebra rubric"

    cell_range = "GRADES!B" + str(count+6) #specific to the gradebook Jana uses (grades tab and B6 onward)

    document = docx.Document(file)
    print(file.name)

    # For debugging

    # print(colored("\n========= Found %d tables in the document ==========" % (len(document.tables)), "blue"))

    # # Print out summary of tables found in the document
    # for t, table in enumerate(document.tables):
    #     print(colored("Table %d has %d rows and %d columns" % (t, len(table.rows), len(table.columns)), "yellow"))

    # # Example usage:
    # result = count_highlights_in_range(document, table_num=0, start_row=3, end_row=3, start_column=1, end_column=3, darkColor=darkColor, lightColor=lightColor)
    # print("Foundational Count:", result[0])
    # print("Proficient Count:", result[1])
    # print("Exemplary Count:", result[2])

    #writing to gradebook

    content = count_highlights_in_range(document, table_num=0, start_row=1, end_row=5, start_column=1, end_column=3, darkColor=darkColor, lightColor=lightColor)

    skills = count_highlights_in_range(document, table_num=1, start_row=1, end_row=9, start_column=1, end_column=3, darkColor=darkColor, lightColor=lightColor)

    habits = count_highlights_in_range(document, table_num=2, start_row=1, end_row=4, start_column=1, end_column=3, darkColor=darkColor, lightColor=lightColor)
    #getting rid of duplicates again and calculating scores for content, habits, skills etc.
    content_f = content[0]
    content_p = content[1]
    content_e = content[2]
    skills_f = skills[0]
    skills_p = skills[1]
    skills_e = skills[2]
    habits_f = habits[0]
    habits_p = habits[1]
    habits_e = habits[2]

    score = [[name,None,None,content_f,content_p,content_e,None,None,skills_f,skills_p,skills_e,None,None,habits_f,habits_p,habits_e]]
    #updating the google sheet
    print("CELL RANGE " + cell_range)
    request = sheet.values().update(spreadsheetId=spreadsheetLink,
                                    range=cell_range,valueInputOption="USER_ENTERED",
                                    body = {"values":score}).execute()



    #writing to nubric gradebook
    all_rows = []
    all_rows_totalStandards = []
    for t, table in enumerate(document.tables):
        row_values = []
        row_values_totalStandards = []
        for r in range(1,len(table.rows)):
            row_values.append(count_highlights_in_range(document, table_num=t, start_row=r, end_row=r, start_column=1, end_column=3, darkColor=darkColor, lightColor=lightColor))
            row_values_totalStandards.append(count_total_in_range(document, table_num = t, start_row = r, end_row = r, start_column = 1, end_column = 3))
        all_rows.append(row_values)
        all_rows_totalStandards.append(row_values_totalStandards)

    #conversion goes here
    input_array = []
   
            
    for i in range(len(all_rows)):
        for x in range(len(all_rows[i])):
            input_value = convertToNubric(all_rows[i][x], all_rows_totalStandards[i][x])

            input_array.append(input_value)

        #BROKEN LINE ABOVE

    null_array = [8,11,14,17,20,22,25,29]

    missed_counter = 0
    internal_count = 0
    print((len(input_array) + len(null_array)))
    print(input_array)
    for i in range(len(input_array) + len(null_array)): # iterate through both the null and input arrays
        letter_index = chr(ord('@')+count+3)


        cell_range_nubric = "GRADEBOOK!" + letter_index + str(internal_count+3)

        if i in null_array:
            request = sheet.values().update(spreadsheetId=spreadsheetNubricLink,
                                            range=cell_range_nubric, valueInputOption="USER_ENTERED",
                                            body={"values": None}).execute()

            missed_counter = missed_counter+1  # missed counter accounts for the None inputs
        else:
            value = [[input_array[i-missed_counter]]]



            request = sheet.values().update(spreadsheetId=spreadsheetNubricLink,
                                            range=cell_range_nubric, valueInputOption="USER_ENTERED",
                                            body={"values": value}).execute()
        
        internal_count +=1    

    count +=1

    #TO DO:
    #go through all rows, get value with function, write as 1-4 to other sheet
