from docx.api import Document
import docx

document = docx.Document('test.docx')
table = document.tables[0]

# Data will be a list of rows represented as dictionaries
# containing each row's data.
data = []


# for paragraph in document.paragraphs:
#     for run in paragraph.runs:
#         if run.font.highlight_color is not None:  #Here I want to add the highlighted condition too
#             data.append("1")
#         print(run.text)
print(len(table.rows))
for i in range(len(table.rows)): #why is it repeating
    print(i)
    for paragraph in table.cell(i,1).paragraphs:
        # print(paragraph.text)
        if paragraph.runs[0].font.highlight_color == 4:
            data.append(paragraph.text)
            #print('green')
    # if table.cell(1,1).paragraphs[0].runs[1].font.highlight_color == 5:
    #     data.append(1)
    #     print('green')

print (list(set(data)))
print (len(list(set(data))))

# keys = None
# for i, row in enumerate(table.rows):
#     text = (cell.text for cell in row.cells) #FIX!!!! bc it's plain text rn

#     # Establish the mapping based on the first row
#     # headers; these will become the keys of our dictionary
#     if i == 0:
#         keys = tuple(text) #FIX!!!! bc it's plain text rn
#         continue

#     # Construct a dictionary for this row, mapping
#     # keys to values for this row
#     row_data = dict(zip(keys, text))

#     if row_data.font.highlight_color is not None:
#         data.append(row_data)

# print(data)