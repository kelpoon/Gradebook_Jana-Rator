from __future__ import print_function
from pydrive.auth import GoogleAuth
from pydrive.drive import GoogleDrive
import os
import os.path
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from google.oauth2 import service_account
from docx.api import Document
import docx
from docx.enum.text import WD_COLOR_INDEX
from termcolor import colored
from pathlib import Path



lightColor = 4
darkColor = 11



SHOW_EVERYTHING_INCLUDING_NON_HIGHLIGHTED = False

# START OF DRIVE THINGS #

MIMETYPES = {
        # Drive Document files as MS dox
        'application/vnd.google-apps.document': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        # Drive Sheets files as MS Excel files.
        'application/vnd.google-apps.spreadsheet': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        # Drive presentation as MS pptx
        'application/vnd.google-apps.presentation': 'application/vnd.openxmlformats-officedocument.presentationml.presentation'
        # see https://developers.google.com/drive/v3/web/mime-types
    }
EXTENSIONS = {
        'application/vnd.google-apps.document': '.docx',
        'application/vnd.google-apps.spreadsheet': '.xlsx',
        'application/vnd.google-apps.presentation': '.pptx'
}

def authenticate():
    gauth = GoogleAuth()

    # Tries to load saved client credentials
    gauth.LoadCredentialsFile("credentials.txt")
    
    if gauth.credentials is None:
        # Authenticate if they're not there
        gauth.LocalWebserverAuth()
    elif gauth.access_token_expired:
        # Refresh them if expired
        gauth.Refresh()
    else:
        # Initialize the saved credentials
        gauth.Authorize()

    # Save the current credentials to a file
    gauth.SaveCredentialsFile("credentials.txt")

    drive = GoogleDrive(gauth)
    return drive

def escape_fname(name):
    return name.replace('/','_')

def search_folder(folder_id, root):
    file_list = drive.ListFile({'q': "'%s' in parents and trashed=false" % folder_id}).GetList()
    for file in file_list:
        # print('title: %s, id: %s, kind: %s' % (file['title'], file['id'], file['mimeType']))
        # print(file)
        if file['mimeType'].split('.')[-1] == 'folder':
            foldername = escape_fname(file['title'])
            create_folder(root,foldername)
            search_folder(file['id'], '{}{}/'.format(root,foldername))
        else:
            download_mimetype = None
            filename = escape_fname(file['title'])
            filename = '{}{}'.format(root,filename)
            try:
                print('DOWNLOADING:', filename)
                if file['mimeType'] in MIMETYPES:
                    download_mimetype = MIMETYPES[file['mimeType']]

                    file.GetContentFile(filename+EXTENSIONS[file['mimeType']], mimetype=download_mimetype)
                else:
                    file.GetContentFile(filename)
            except:
                print('FAILED')
                f.write(filename+'\n')

def create_folder(path,name):
    os.mkdir('{}{}'.format(path,escape_fname(name)))

if __name__ == '__main__':
    drive = authenticate()

    f = open("failed.txt","w+")
    folder_id = '1tyeoX2ZgmgSW57cS8z-r4S_Z6hKPl2yo'
    root = 'drive_download'
    if not os.path.exists(root):
        os.makedirs(root)

    search_folder(folder_id,root+'/')
    f.close() 


# END OF DRIVE THINGS #

dir_list = os.listdir(root)
print(dir_list)


# If modifying these scopes, delete the file token.json.

# The ID and range of a sample spreadsheet.
SAMPLE_SPREADSHEET_ID = '1-CZJljjGtjP9J2AX5n80jlOdPWBCvum5pm7F8rDS510'


        # Call the Sheets API
#result = sheet.values().get(spreadsheetId=SAMPLE_SPREADSHEET_ID,
#                                   range="GRADES!A1:G13").execute()
#values = result.get('values', [])

SCOPES = ['https://www.googleapis.com/auth/spreadsheets']
SERVICE_ACCOUNT_FILE = 'keys.json'
creds = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE, scopes=SCOPES)
service = build('sheets', 'v4', credentials=creds)
sheet = service.spreadsheets()


# for outputting pretty colors to the terminal
def convert_wd_color_index_to_termcolor(color_index):
    if (color_index == WD_COLOR_INDEX.BLUE):
        return "blue"
    if (color_index == WD_COLOR_INDEX.TURQUOISE):
        return "cyan"
    if (color_index == WD_COLOR_INDEX.GREEN):
        return "green"
    if (color_index == WD_COLOR_INDEX.BRIGHT_GREEN):
        return "light_green"
    if (color_index == WD_COLOR_INDEX.RED):
        return "red"
    if (color_index == WD_COLOR_INDEX.YELLOW):
        return "yellow"

    print("WARNING: Unrecognized color index: %s" % (color_index))
    return "white"


def calculateScoreFromHighlights(highlights):
    score = 0
    for h in highlights:
        score += h[1]
    return score

count = 0


for file in Path(root).iterdir():
    name = "test " + str(count)
    print(name)

    cell_range = "GRADES!B" + str(count+6)

    print(cell_range)
    
    document = docx.Document(file)

    print(colored("\n========= Found %d tables in the document ==========" % (len(document.tables)), "blue"))

    # Print out summary of tables found in the document
    for t, table in enumerate(document.tables):
        print(colored("Table %d has %d rows and %d columns" % (t, len(table.rows), len(table.columns)), "yellow"))

    # each of these lists will contain tuples of (text, score) which we'll later remove dupes using set


    allFoundational = []
    allProficients = []
    allExemplarys = []


    # Print out detailed contents of each table, along with what is highlighted
    for t, table in enumerate(document.tables):
        highlightedFoundationals = []
        highlightedProficients = []
        highlightedExemplarys = []
        print("\n\nTABLE %d:" % (t))
        for r, row in enumerate(table.rows):
            print("\n-------- Row %d --------" % (r))
            for c, cell in enumerate(row.cells):
                print()
                for p, paragraph in enumerate(cell.paragraphs):
                    if (len(paragraph.runs) == 0):
                        continue
                    # score will be the precentage of runs inside this paragraph that are highlighted
                    numHighlightedRuns = 0
                    text = paragraph.text
                    for r2, run in enumerate(paragraph.runs):
                        colors_foundational = []
                        colors_proficient = []
                        colors_exemplary = []
                        if (c == 1):
                            for i in range(len(paragraph.runs)):
                                colors_foundational.append(paragraph.runs[i].font.highlight_color)  
                                                
                        if (c == 2):
                            for i in range(len(paragraph.runs)):
                                colors_proficient.append(paragraph.runs[i].font.highlight_color)

                        if (c == 3):
                            for i in range(len(paragraph.runs)):
                                colors_exemplary.append(paragraph.runs[i].font.highlight_color)
                        colors_foundational = list(set(colors_foundational))
                        colors_proficient = list(set(colors_proficient))
                        colors_exemplary = list(set(colors_exemplary))
                        
                        
                        if len(colors_foundational) == 1:
                            if colors_foundational[0] == darkColor:
                                highlightedFoundationals.append((text,1))
                            if colors_foundational[0] == lightColor:
                                highlightedFoundationals.append((text,.5))
                            
                        elif len(colors_foundational) > 1:
                            if darkColor and lightColor in colors_foundational:
                                highlightedFoundationals.append((text,.75))
                            elif darkColor in colors_foundational:
                                highlightedFoundationals.append((text,.5))
                            elif lightColor in colors_foundational:
                                highlightedFoundationals.append([text,.25])
                        

                        if len(colors_proficient) == 1:
                            if colors_proficient[0] == darkColor:
                                highlightedProficients.append((text,1))
                            if colors_proficient[0] == lightColor:
                                highlightedProficients.append((text,.5))
                            # elif colors_proficient[0]!=None:
                            #     print(colors_proficient)

                            
                        elif len(colors_proficient) > 1:
                            if darkColor and lightColor in colors_proficient:
                                highlightedProficients.append((text,.75))
                            elif darkColor in colors_proficient:
                                highlightedProficients.append((text,.5))
                            elif lightColor in colors_proficient:
                                highlightedProficients.append((text,.25))

                        if len(colors_exemplary) == 1:
                            if colors_exemplary[0] == darkColor:
                                highlightedExemplarys.append((text,1))
                            if colors_exemplary[0] == lightColor:
                                highlightedExemplarys.append((text,.5))

                            
                        elif len(colors_exemplary) > 1:
                            if darkColor and lightColor in colors_exemplary:
                                highlightedExemplarys.append((text,.75))
                            elif darkColor in colors_exemplary:
                                highlightedExemplarys.append((text,.5))
                            elif lightColor in colors_exemplary:
                                highlightedExemplarys.append((text,.25))
                    
                    for r2, run in enumerate(paragraph.runs):
                        if run.font.highlight_color is not None:
                            print(
                                colored("*Table %d, Row %d, Cell %d, Paragraph %d, Run %d: %s" % (t, r, c, p, r2, run.text),
                                        convert_wd_color_index_to_termcolor(run.font.highlight_color)))
                        else:
                            if SHOW_EVERYTHING_INCLUDING_NON_HIGHLIGHTED:
                                print(" Table %d, Row %d, Cell %d, Paragraph %d, Run %d: %s" % (t, r, c, p, r2, run.text))
        allFoundational.append(highlightedFoundationals)
        allProficients.append(highlightedProficients)
        allExemplarys.append(highlightedExemplarys)

    print("\n\n========= FOUNDATIONALS ==========")
    # highlightedFoundationals = list(set(highlightedFoundationals))
    # print(highlightedFoundationals)
    # print(calculateScoreFromHighlights(highlightedFoundationals))
    for i in allFoundational:
        print(calculateScoreFromHighlights(list(set(i))))




    print("\n\n========= PROFICIENTS ==========")
    # highlightedProficients = list(set(highlightedProficients))
    # # print(highlightedProficients)
    # print(calculateScoreFromHighlights(highlightedProficients))
    for i in allProficients:
        print(calculateScoreFromHighlights(list(set(i))))

    print("\n\n========= EXEMPLARYS ==========")
    # highlightedExemplarys = list(set(highlightedExemplarys))
    # # print(highlightedExemplarys)
    # print(calculateScoreFromHighlights(highlightedExemplarys))
    for i in allExemplarys:
        print(calculateScoreFromHighlights(list(set(i))))


    content_f = calculateScoreFromHighlights(list(set(allFoundational[0])))
    content_p = calculateScoreFromHighlights(list(set(allProficients[0])))
    content_e = calculateScoreFromHighlights(list(set(allExemplarys[0])))
    skills_f = calculateScoreFromHighlights(list(set(allFoundational[1])))
    skills_p = calculateScoreFromHighlights(list(set(allProficients[1])))
    skills_e = calculateScoreFromHighlights(list(set(allExemplarys[1])))
    habits_f = calculateScoreFromHighlights(list(set(allFoundational[2])))
    habits_p = calculateScoreFromHighlights(list(set(allProficients[2])))
    habits_e = calculateScoreFromHighlights(list(set(allExemplarys[2])))

    score = [[name,None,None,content_f,content_p,content_e,None,None,skills_f,skills_p,skills_e,None,None,habits_f,habits_p,habits_e]]
    aoa = [["1/1/2020",4000]]
    request = sheet.values().update(spreadsheetId=SAMPLE_SPREADSHEET_ID,
                                    range=cell_range,valueInputOption="USER_ENTERED",
                                    body = {"values":score}).execute()
    count +=1
