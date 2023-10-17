from docx.api import Document
import docx
from docx.enum.text import WD_COLOR_INDEX
from termcolor import colored

SHOW_EVERYTHING_INCLUDING_NON_HIGHLIGHTED = False

# for outputting pretty colors to the terminal
def convert_wd_color_index_to_termcolor(color_index):
    if (color_index == WD_COLOR_INDEX.BLUE):
        return "blue"
    if (color_index == WD_COLOR_INDEX.TURQUOISE):
        return "cyan"
    if (color_index == WD_COLOR_INDEX.GREEN): #11
        return "green"
    if (color_index == WD_COLOR_INDEX.BRIGHT_GREEN): #4
        return "light_green"
    if (color_index == WD_COLOR_INDEX.RED):
        return "red"
    if (color_index == WD_COLOR_INDEX.YELLOW):
        return "yellow"

    print("WARNING: Unrecognized color index: %s" % (color_index))
    return "white"

def calculateScoreFromHighlights(highlights):
    score = 0
    for h in highlights:
        score += h[1]
    return score

document = docx.Document('test2.docx')

print(colored("\n========= Found %d tables in the document ==========" % (len(document.tables)), "blue"))

# Print out summary of tables found in the document
for t, table in enumerate(document.tables):
    print(colored("Table %d has %d rows and %d columns" % (t, len(table.rows), len(table.columns)), "yellow"))    

# each of these lists will contain tuples of (text, score) which we'll later remove dupes using set
highlightedFoundationals = []
highlightedProficients = []
highlightedExemplarys = []

# Print out detailed contents of each table, along with what is highlighted
for t, table in enumerate(document.tables):
    print("\n\nTABLE %d:" % (t))
    for r, row in enumerate(table.rows):
        print("\n-------- Row %d --------" % (r))
        for c, cell in enumerate(row.cells):
            print()
            for p, paragraph in enumerate(cell.paragraphs):
                if (len(paragraph.runs) == 0):
                    continue
                # score will be the precentage of runs inside this paragraph that are highlighted
                numHighlightedRuns = 0
                text = paragraph.text
                for r2, run in enumerate(paragraph.runs):
                    colors_foundational = []
                    colors_proficient = []
                    colors_exemplary = []
                    if (c == 1):
                        for i in range(len(paragraph.runs)):
                            colors_foundational.append(paragraph.runs[i].font.highlight_color)  
                                            
                    if (c == 2):
                        for i in range(len(paragraph.runs)):
                            colors_proficient.append(paragraph.runs[i].font.highlight_color)

                    if (c == 3):
                        for i in range(len(paragraph.runs)):
                            colors_exemplary.append(paragraph.runs[i].font.highlight_color)
                    colors_foundational = list(set(colors_foundational))
                    colors_proficient = list(set(colors_proficient))
                    colors_exemplary = list(set(colors_exemplary))
                    
                    
                    if len(colors_foundational) == 1:
                        if colors_foundational[0] == WD_COLOR_INDEX.GREEN:
                            highlightedFoundationals.append((text,1))
                        if colors_foundational[0] == WD_COLOR_INDEX.BRIGHT_GREEN:
                            highlightedFoundationals.append((text,.5))
                        
                    elif len(colors_foundational) > 1:
                        if WD_COLOR_INDEX.GREEN and WD_COLOR_INDEX.BRIGHT_GREEN in colors_foundational:
                            highlightedFoundationals.append((text,.75))
                        elif WD_COLOR_INDEX.GREEN in colors_foundational:
                            highlightedFoundationals.append((text,.5))
                        elif WD_COLOR_INDEX.BRIGHT_GREEN in colors_foundational:
                            highlightedFoundationals.append([text,.25])
                    

                    if len(colors_proficient) == 1:
                        if colors_proficient[0] == WD_COLOR_INDEX.GREEN:
                            highlightedProficients.append((text,1))
                        if colors_proficient[0] == WD_COLOR_INDEX.BRIGHT_GREEN:
                            highlightedProficients.append((text,.5))
                        # elif colors_proficient[0]!=None:
                        #     print(colors_proficient)

                        
                    elif len(colors_proficient) > 1:
                        if WD_COLOR_INDEX.GREEN and WD_COLOR_INDEX.BRIGHT_GREEN in colors_proficient:
                            highlightedProficients.append((text,.75))
                        elif WD_COLOR_INDEX.GREEN in colors_proficient:
                            highlightedProficients.append((text,.5))
                        elif WD_COLOR_INDEX.BRIGHT_GREEN in colors_proficient:
                            highlightedProficients.append((text,.25))

                    if len(colors_exemplary) == 1:
                        if colors_exemplary[0] == WD_COLOR_INDEX.GREEN:
                            highlightedExemplarys.append((text,1))
                        if colors_exemplary[0] == WD_COLOR_INDEX.BRIGHT_GREEN:
                            highlightedExemplarys.append((text,.5))

                        
                    elif len(colors_exemplary) > 1:
                        if WD_COLOR_INDEX.GREEN and WD_COLOR_INDEX.BRIGHT_GREEN in colors_exemplary:
                            highlightedExemplarys.append((text,.75))
                        elif WD_COLOR_INDEX.GREEN in colors_exemplary:
                            highlightedExemplarys.append((text,.5))
                        elif WD_COLOR_INDEX.BRIGHT_GREEN in colors_exemplary:
                            highlightedExemplarys.append((text,.25))
                

                for r2, run in enumerate(paragraph.runs):
                    if run.font.highlight_color is not None:
                        print(colored("*Table %d, Row %d, Cell %d, Paragraph %d, Run %d: %s" % (t, r, c, p, r2, run.text), convert_wd_color_index_to_termcolor(run.font.highlight_color)))
                    else:
                        if SHOW_EVERYTHING_INCLUDING_NON_HIGHLIGHTED:
                            print(" Table %d, Row %d, Cell %d, Paragraph %d, Run %d: %s" % (t, r, c, p, r2, run.text))


print("\n\n========= FOUNDATIONALS ==========")
highlightedFoundationals = list(set(highlightedFoundationals))
print(highlightedFoundationals)
print(calculateScoreFromHighlights(highlightedFoundationals))

print("\n\n========= PROFICIENTS ==========")
highlightedProficients = list(set(highlightedProficients))
# print(highlightedProficients)
print(calculateScoreFromHighlights(highlightedProficients))

print("\n\n========= EXEMPLARYS ==========")
highlightedExemplarys = list(set(highlightedExemplarys))
# print(highlightedExemplarys)
print(calculateScoreFromHighlights(highlightedExemplarys))

